"""
PDF Operations Backend for SafePDF
Implements various PDF manipulation operations using PyPDF2/pypdf and Pillow
"""

import os
from io import BytesIO
from os import path as os_path
from pathlib import Path
from tempfile import mkstemp as tmp_mkstemp
from typing import List, Tuple

from SafePDF.logger.logging_config import get_logger

try:
    import tkinter as tk
    from tkinter import Button, Label, Toplevel, messagebox
except ImportError:
    messagebox = Toplevel = Label = Button = tk = None

try:
    from PyPDF2 import PdfReader, PdfWriter
    from PyPDF2.errors import PdfReadError
except ImportError:
    print("Warning: PyPDF2 not installed. PDF operations will not work.")
    PdfReader = PdfWriter = None

try:
    import fitz  # PyMuPDF for better PDF to image conversion
    from PIL import Image, ImageTk
except ImportError:
    print("Warning: PIL/Pillow or PyMuPDF not installed. Some operations may not work.")
    Image = ImageTk = fitz = None

class PDFOperations:
    """Class containing all PDF manipulation operations"""
    
    def __init__(self, progress_callback=None):
        """
        Initialize PDF operations handler
        
        Args:
            progress_callback: Function to call for progress updates (0-100)
        """
        self.progress_callback = progress_callback
        # Cancellation flag that can be set by controller/UI
        self._cancel_requested = False

        # Module logger
        self.logger = get_logger('SafePDF.PDFOps')
        
    def _ensure_parent_dir(self, file_path: str):
        """
        Ensure the parent directory for `file_path` exists. If `file_path`
        does not contain a directory component, do nothing.
        """
        try:
            parent = os_path.dirname(file_path) if file_path else ''
            if parent:
                os.makedirs(parent, exist_ok=True)
        except Exception:
            # Creation failure should be handled by the caller when opening/writing files.
            self.logger.error(f"Failed to create parent directory for {file_path}", exc_info=True)
            pass

    def _atomic_write_file(self, final_path: str, write_func):
        """
        Atomically write to `final_path` using a temp file in the same directory.
        `write_func` is called with an open file object (binary mode) to write content.
        Uses os.replace for atomic move and ensures cleanup on errors.
        """
        parent = os_path.dirname(final_path) or os.getcwd()
        self._ensure_parent_dir(final_path)
        
        fd = None
        tmp_path = None
        try:
            # Create secure temp file in same directory as target
            fd, tmp_path = tmp_mkstemp(
                prefix=".safepdf_tmp_",
                suffix=os_path.splitext(final_path)[1] or ".tmp",
                dir=parent
            )
            
            # Write content via callback
            with os.fdopen(fd, "wb") as tmpf:
                fd = None  # fdopen takes ownership
                write_func(tmpf)
                tmpf.flush()
                try:
                    os.fsync(tmpf.fileno())
                except (OSError, AttributeError):
                    pass
            
            # Atomic replace
            os.replace(tmp_path, final_path)
            tmp_path = None  # Successfully moved
            
        except Exception:
            self.logger.error(f"Error during atomic write to {final_path}", exc_info=True)
            raise
        finally:
            # Cleanup on error
            if fd is not None:
                try:
                    os.close(fd)
                except Exception:
                    self.logger.error("Error closing file descriptor", exc_info=True)
                    pass
            if tmp_path and os_path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    self.logger.error("Error removing temporary file", exc_info=True)
                    pass

    def _atomic_write_via_path(self, final_path: str, write_path_func):
        """
        Atomically write to `final_path` by providing a temp path to write_path_func.
        `write_path_func` is called with a temp file path (string) to write to.
        Uses os.replace for atomic move and ensures cleanup on errors.
        """
        parent = os_path.dirname(final_path) or os.getcwd()
        self._ensure_parent_dir(final_path)
        
        fd = None
        tmp_path = None
        try:
            # Create secure temp file in same directory
            fd, tmp_path = tmp_mkstemp(
                prefix=".safepdf_tmp_",
                suffix=os_path.splitext(final_path)[1] or ".tmp",
                dir=parent
            )
            os.close(fd)
            fd = None
            
            # Let caller write to temp path
            write_path_func(tmp_path)
            
            # Atomic replace
            os.replace(tmp_path, final_path)
            tmp_path = None  # Successfully moved
            
        except Exception:
            raise
        finally:
            # Cleanup on error
            if fd is not None:
                try:
                    os.close(fd)
                except Exception:
                    pass
            if tmp_path and os_path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    def update_progress(self, value):
        """Update progress if callback is available"""
        if self.progress_callback:
            self.progress_callback(value)

    def request_cancel(self):
        """Request cancellation of a running operation."""
        self._cancel_requested = True
            
    def validate_pdf(self, file_path: str) -> bool:
        """
        Validate if file is a valid PDF
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            True if valid PDF, False otherwise
        """
        try:
            if not PdfReader:
                return False
                
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                # Try to access pages to ensure it's readable
                len(reader.pages)
            return True
        except (PdfReadError, Exception):
            return False
            
    def compress_pdf(self, input_path: str, output_path: str, quality: str = "medium") -> Tuple[bool, str]:
        """
        Compress PDF file
        
        Args:
            input_path: Input PDF file path
            output_path: Output PDF file path
            quality: Compression quality ("low", "medium", "high")
            
        Returns:
            Tuple of (success, message)
        """
        try:
            if not PdfReader or not PdfWriter:
                return False, "PyPDF2/pypdf not available"
            
            self.update_progress(5)
            
            # Validate input file
            if not os_path.exists(input_path):
                return False, "Input file does not exist"
                
            if not self.validate_pdf(input_path):
                return False, "Input file is not a valid PDF"
            
            # Prefer PyMuPDF approach for effective compression (image re-encoding)
            if fitz:
                # Map quality to dpi and jpeg quality
                if quality == "low":
                    dpi = 100
                    jpeg_q = 40
                elif quality == "high":
                    dpi = 220
                    jpeg_q = 85
                elif quality == "ultra":
                    dpi = 300
                    jpeg_q = 95
                else:  # medium
                    dpi = 150
                    jpeg_q = 60
                
                doc = fitz.open(input_path)
                total_pages = len(doc)
                if total_pages == 0:
                    return False, "Input PDF has no pages"
                
                self.update_progress(15)
                
                new_doc = fitz.open()  # empty document to populate with images
                for i in range(total_pages):
                    # Check for cancellation request
                    if self._cancel_requested:
                        try:
                            new_doc.close()
                            doc.close()
                        except Exception:
                            pass
                        return False, "Operation cancelled by user"
                    page = doc.load_page(i)
                    # render page at chosen dpi
                    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    
                    # Try to produce JPEG bytes for good compression
                    img_bytes = None
                    if Image is not None:
                        try:
                            mode = "RGB" if pix.n < 4 else "RGBA"
                            img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                            buf = BytesIO()
                            img.save(buf, format="JPEG", quality=jpeg_q, optimize=True)
                            img_bytes = buf.getvalue()
                        except Exception:
                            img_bytes = None
                    # Fallback to Pixmap's PNG bytes if PIL not available or failed
                    if img_bytes is None:
                        try:
                            # PyMuPDF provides getPNGData() for pixmap
                            img_bytes = pix.getPNGData()
                        except Exception:
                            # As last resort, use raw pixmap bytes (may be large)
                            try:
                                img_bytes = pix.tobytes()
                            except Exception:
                                img_bytes = None
                    
                    # compute page size in points (1 point = 1/72 inch)
                    width_pts = (pix.width * 72.0) / dpi
                    height_pts = (pix.height * 72.0) / dpi
                    
                    page_rect = fitz.Rect(0, 0, width_pts, height_pts)
                    new_page = new_doc.new_page(width=width_pts, height=height_pts)
                    
                    if img_bytes:
                        try:
                            new_page.insert_image(page_rect, stream=img_bytes, keep_proportion=True)
                        except Exception:
                            # if inserting stream fails, try writing a temp image file and insert by filename
                            tmp_fd = None
                            tmp_name = None
                            try:
                                tmp_fd, tmp_name = tmp_mkstemp(suffix=".avif")
                                os.write(tmp_fd, img_bytes)
                                os.close(tmp_fd)
                                tmp_fd = None
                                new_page.insert_image(page_rect, filename=tmp_name, keep_proportion=True)
                            except Exception:
                                # if even that fails, skip page (should be rare)
                                pass
                            finally:
                                if tmp_fd is not None:
                                    try:
                                        os.close(tmp_fd)
                                    except Exception:
                                        pass
                                if tmp_name and os_path.exists(tmp_name):
                                    try:
                                        os.remove(tmp_name)
                                    except Exception:
                                        pass
                    
                    self.update_progress(15 + (75 * i // total_pages))
                
                # Save new PDF atomically using temp file
                def _save_compressed(tmp_path):
                    try:
                        new_doc.save(tmp_path, deflate=True, garbage=4)
                    except TypeError:
                        # Older PyMuPDF may not accept those kwargs
                        new_doc.save(tmp_path)
                
                try:
                    self._atomic_write_via_path(output_path, _save_compressed)
                finally:
                    new_doc.close()
                    doc.close()
                
                self.update_progress(100)
                
                # verify and compare sizes
                if os_path.exists(output_path) and self.validate_pdf(output_path):
                    original_size = os_path.getsize(input_path)
                    compressed_size = os_path.getsize(output_path)
                    
                    # Guard against zero-size original file
                    if original_size == 0:
                        return False, "Original file size is zero. Cannot calculate compression."
                    
                    if compressed_size < original_size:
                        compression_ratio = (1 - (compressed_size / original_size)) * 100
                        return True, f"PDF compressed successfully. Quality: {quality}. Size reduced by {abs(compression_ratio):.1f}%"
                    else:
                        # If compression didn't reduce size, try with lower quality automatically
                        if quality == "high":
                            # Try medium quality
                            return self.compress_pdf(input_path, output_path, "medium")
                        elif quality == "medium":
                            # Try low quality
                            return self.compress_pdf(input_path, output_path, "low")
                        else:
                            # Already tried low quality, show warning
                            self._show_compression_error_popup()
                            if compressed_size == original_size:
                                return False, "No size reduction achieved. Please try a different quality setting or use 'Microsoft Print to PDF' from the print dialog."
                            else:
                                increase_pct = ((compressed_size / original_size) - 1) * 100
                                return False, f"Compression increased file size by {increase_pct:.1f}%. Please try a different quality setting."
                
                return False, "Compression completed but output file is invalid"
            
            # Fallback: attempt in-place stream compression using PdfWriter (may not always reduce size)
            # The existing writer-based approach is retained as fallback to avoid removing functionality.
            # Minimal fallback implementation:
            self.update_progress(10)
            reader = PdfReader(input_path)
            writer = PdfWriter()
            total_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                if self._cancel_requested:
                    return False, "Operation cancelled by user"
                try:
                    page.compress_content_streams()
                except Exception:
                    pass
                writer.add_page(page)
                self.update_progress(10 + (80 * i // max(1, total_pages)))
            
            def _write_compressed(tmpf):
                writer.write(tmpf)
            
            self._atomic_write_file(output_path, _write_compressed)
            self.update_progress(100)
            
            # Compare sizes and warn if increased
            if os_path.exists(output_path) and self.validate_pdf(output_path):
                original_size = os_path.getsize(input_path)
                compressed_size = os_path.getsize(output_path)
                if original_size == 0:
                    return False, "Original file size is zero. Cannot calculate compression."
                if compressed_size < original_size:
                    compression_ratio = (1 - (compressed_size / original_size)) * 100
                    return True, f"PDF compressed successfully (fallback). Quality: {quality}. Size reduced by {abs(compression_ratio):.1f}%"
                else:
                    # Show warning immediately when compression doesn't reduce size
                    self._show_compression_error_popup()
                    if compressed_size == original_size:
                        return False, "No size reduction achieved using fallback method. Please try a different quality setting or use 'Microsoft Print to PDF' from the print dialog."
                    else:
                        increase_pct = ((compressed_size / original_size) - 1) * 100
                        return False, f"Fallback compression increased file size by {increase_pct:.1f}%. Please try a different quality setting."
            
            return False, "Fallback compression completed but output file is invalid"
            
        except Exception as e:
            return False, f"Compression failed: {str(e)}"
            
    def split_pdf(self, input_path: str, output_dir: str, method: str = "pages", page_range: str = None) -> Tuple[bool, str]:
        """
        Split PDF into multiple files
        
        Args:
            input_path: Input PDF file path
            output_dir: Directory to save split files
            method: Split method ("pages" for each page, "range" for specific range)
            page_range: Page range if method is "range" (e.g., "1-5,7,10-12")
            
        Returns:
            Tuple of (success, message)
        """
        try:
            if not PdfReader or not PdfWriter:
                return False, "PyPDF2/pypdf not available"
                
            self.update_progress(10)
            
            with open(input_path, 'rb') as input_file:
                reader = PdfReader(input_file)
                total_pages = len(reader.pages)
                
                self.update_progress(20)
                
                if method == "pages":
                    # Split each page into separate file
                    for i, page in enumerate(reader.pages):
                        if self._cancel_requested:
                            return False, "Operation cancelled by user"
                        writer = PdfWriter()
                        writer.add_page(page)
                        
                        output_filename = f"page_{i+1}.pdf"
                        output_path = os_path.join(output_dir, output_filename)
                        
                        def _write_page(tmpf):
                            writer.write(tmpf)
                        
                        self._atomic_write_file(output_path, _write_page)
                        self.update_progress(20 + (70 * i // total_pages))
                        
                    self.update_progress(100)
                    return True, f"PDF split into {total_pages} files"
                    
                elif method == "range" and page_range:
                    # Parse page range and create files
                    ranges = self._parse_page_range(page_range, total_pages)
                    
                    for i, (start, end) in enumerate(ranges):
                        if self._cancel_requested:
                            return False, "Operation cancelled by user"
                        writer = PdfWriter()
                        for page_num in range(start-1, end):
                            if 0 <= page_num < total_pages:
                                writer.add_page(reader.pages[page_num])
                        
                        output_filename = f"pages_{start}-{end}.pdf"
                        output_path = os_path.join(output_dir, output_filename)
                        
                        def _write_range(tmpf):
                            writer.write(tmpf)
                        
                        self._atomic_write_file(output_path, _write_range)
                        self.update_progress(20 + (70 * i // len(ranges)))
                        
                    self.update_progress(100)
                    return True, f"PDF split into {len(ranges)} files based on ranges"
                    
            return False, "Invalid split method or parameters"
            
        except Exception as e:
            return False, f"Split failed: {str(e)}"
            
    def merge_pdfs(self, input_paths: List[str], output_path: str) -> Tuple[bool, str]:
        """
        Merge multiple PDF files
        
        Args:
            input_paths: List of input PDF file paths
            output_path: Output merged PDF file path
            
        Returns:
            Tuple of (success, message)
        """
        try:
            if not PdfReader or not PdfWriter:
                return False, "PyPDF2/pypdf not available"
                
            self.update_progress(10)
            
            writer = PdfWriter()
            total_files = len(input_paths)
            
            for i, input_path in enumerate(input_paths):
                if self._cancel_requested:
                    return False, "Operation cancelled by user"
                with open(input_path, 'rb') as input_file:
                    reader = PdfReader(input_file)
                    for page in reader.pages:
                        writer.add_page(page)
                        
                self.update_progress(10 + (80 * i // total_files))
            
            def _write_merged(tmpf):
                writer.write(tmpf)
            
            self._atomic_write_file(output_path, _write_merged)
            self.update_progress(100)
            return True, f"Successfully merged {total_files} PDF files"
            
        except Exception as e:
            return False, f"Merge failed: {str(e)}"
            
    def pdf_to_jpg(self, input_path: str, output_dir: str, dpi: int = 200) -> Tuple[bool, str]:
        """
        Convert PDF pages to JPG images
        
        Args:
            input_path: Input PDF file path
            output_dir: Directory to save JPG files
            dpi: Resolution for images
            
        Returns:
            Tuple of (success, message)
        """
        try:
            if not fitz:
                return False, "PyMuPDF not available for PDF to image conversion"
                
            self.update_progress(10)
            
            # Open PDF
            pdf_document = fitz.open(input_path)
            total_pages = len(pdf_document)
            
            self.update_progress(20)
            
            for page_num in range(total_pages):
                if self._cancel_requested:
                    return False, "Operation cancelled by user"
                page = pdf_document.load_page(page_num)
                
                # Create transformation matrix for desired DPI
                mat = fitz.Matrix(dpi/72, dpi/72)
                pix = page.get_pixmap(matrix=mat)
                
                # Save as JPG
                output_filename = f"page_{page_num + 1}.avif"
                output_path = os_path.join(output_dir, output_filename)
                pix.save(output_path)
                
                self.update_progress(20 + (70 * page_num // total_pages))
            
            pdf_document.close()
            self.update_progress(100)
            
            return True, f"Converted {total_pages} pages to JPG images"
            
        except Exception as e:
            return False, f"PDF to JPG conversion failed: {str(e)}"
            
    def rotate_pdf(self, input_path: str, output_path: str, angle: int = 90) -> Tuple[bool, str]:
        """
        Rotate PDF pages
        
        Args:
            input_path: Input PDF file path
            output_path: Output PDF file path
            angle: Rotation angle (90, 180, 270)
            
        Returns:
            Tuple of (success, message)
        """
        try:
            if not PdfReader or not PdfWriter:
                return False, "PyPDF2/pypdf not available"
                
            self.update_progress(10)
            
            with open(input_path, 'rb') as input_file:
                reader = PdfReader(input_file)
                writer = PdfWriter()
                
                total_pages = len(reader.pages)
                self.update_progress(30)
                
                for i, page in enumerate(reader.pages):
                    if self._cancel_requested:
                        return False, "Operation cancelled by user"
                    rotated_page = page.rotate(angle)
                    writer.add_page(rotated_page)
                    self.update_progress(30 + (60 * i // total_pages))
                
                def _write_rotated(tmpf):
                    writer.write(tmpf)
                
                self._atomic_write_file(output_path, _write_rotated)
            self.update_progress(100)
            return True, f"PDF rotated by {angle} degrees"
            
        except Exception as e:
            return False, f"Rotation failed: {str(e)}"
            
    def repair_pdf(self, input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Attempt to repair a corrupted PDF
        
        Args:
            input_path: Input PDF file path
            output_path: Output repaired PDF file path
            
        Returns:
            Tuple of (success, message)
        """
        try:
            if not PdfReader or not PdfWriter:
                return False, "PyPDF2/pypdf not available"
                
            self.update_progress(10)
            
            with open(input_path, 'rb') as input_file:
                reader = PdfReader(input_file, strict=False)  # Less strict parsing
                writer = PdfWriter()
                
                self.update_progress(30)
                
                pages_recovered = 0
                total_pages = len(reader.pages) if hasattr(reader, 'pages') else 0
                
                try:
                    for i, page in enumerate(reader.pages):
                        if self._cancel_requested:
                            return False, "Operation cancelled by user"
                        try:
                            writer.add_page(page)
                            pages_recovered += 1
                            if total_pages > 0:
                                self.update_progress(30 + (60 * i // total_pages))
                        except Exception:
                            # Skip corrupted pages
                            continue
                            
                except Exception:
                    pass  # Continue with whatever pages we could recover
                
                if pages_recovered > 0:
                    def _write_repaired(tmpf):
                        writer.write(tmpf)
                    
                    self._atomic_write_file(output_path, _write_repaired)
                    self.update_progress(100)
                    return True, f"PDF repaired. Recovered {pages_recovered} pages"
                else:
                    return False, "Could not recover any pages from the PDF"
                    
        except Exception as e:
            return False, f"Repair failed: {str(e)}"
            
    def _parse_page_range(self, page_range: str, total_pages: int) -> List[Tuple[int, int]]:
        """
        Parse page range string into list of (start, end) tuples
        
        Args:
            page_range: Range string like "1-5,7,10-12"
            total_pages: Total number of pages in PDF
            
        Returns:
            List of (start, end) tuples
        """
        ranges = []
        parts = page_range.split(',')
        
        for part in parts:
            part = part.strip()
            if '-' in part:
                start, end = part.split('-', 1)
                start = max(1, min(int(start.strip()), total_pages))
                end = max(start, min(int(end.strip()), total_pages))
                ranges.append((start, end))
            else:
                page_num = max(1, min(int(part.strip()), total_pages))
                ranges.append((page_num, page_num))
                
        return ranges
        
    def get_pdf_info(self, file_path: str) -> dict:
        """
        Get basic information about a PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with PDF information
        """
        try:
            if not PdfReader:
                return {"error": "PyPDF2/pypdf not available"}
                
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                info = {
                    "pages": len(reader.pages),
                    "file_size": os_path.getsize(file_path),
                    "file_name": os_path.basename(file_path)
                }
                
                if reader.metadata:
                    info.update({
                        "title": reader.metadata.get('/Title', 'Unknown'),
                        "author": reader.metadata.get('/Author', 'Unknown'),
                        "creator": reader.metadata.get('/Creator', 'Unknown'),
                        "producer": reader.metadata.get('/Producer', 'Unknown'),
                    })
                    
                return info
                
        except Exception as e:
            return {"error": str(e)}
    
    def pdf_to_txt(self, input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Extract text from PDF and save to TXT file
        
        Args:
            input_path: Path to input PDF file
            output_path: Path to output TXT file
            
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            if not PdfReader:
                return False, "PyPDF2/pypdf not available"
            
            with open(input_path, 'rb') as file:
                reader = PdfReader(file)
                text_content = ""
                
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    self.update_progress(int((i + 1) / total_pages * 100))
                    if self._cancel_requested:
                        return False, "Operation cancelled"
                    
                    text_content += page.extract_text() + "\n\n"
                
            def _write_text(tmpf):
                tmpf.write(text_content.encode('utf-8'))
            
            self._atomic_write_file(output_path, _write_text)
            return True, f"Text extracted to {output_path}"
            
        except Exception as e:
            return False, f"Text extraction failed: {str(e)}"
    
    def extract_hidden_info(self, input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Extract hidden information and metadata from PDF
        
        Args:
            input_path: Path to input PDF file
            output_path: Path to output TXT file with extracted info
            
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            if not PdfReader:
                return False, "PyPDF2/pypdf not available"
            
            info = self.get_pdf_info(input_path)
            if "error" in info:
                return False, info["error"]
            
            # Extract additional hidden information
            hidden_info = []
            hidden_info.append("=== PDF METADATA ===")
            hidden_info.append(f"Title: {info.get('title', 'N/A')}")
            hidden_info.append(f"Author: {info.get('author', 'N/A')}")
            hidden_info.append(f"Creator: {info.get('creator', 'N/A')}")
            hidden_info.append(f"Producer: {info.get('producer', 'N/A')}")
            hidden_info.append(f"Pages: {info.get('pages', 'N/A')}")
            hidden_info.append(f"File Size: {info.get('file_size', 'N/A')} bytes")
            
            # Try to extract more detailed info
            try:
                with open(input_path, 'rb') as file:
                    reader = PdfReader(file)
                    if reader.trailer and "/Info" in reader.trailer:
                        info_dict = reader.trailer["/Info"]
                        hidden_info.append("\n=== ADDITIONAL INFO DICTIONARY ===")
                        for key, value in info_dict.items():
                            hidden_info.append(f"{key}: {value}")
            except Exception as e:
                hidden_info.append(f"\nError extracting additional info: {str(e)}")
            
            hidden_info.append("\n=== END OF EXTRACTED INFORMATION ===")
            hidden_info.append("These details are extracted by SafePDF.")
            
            # Write to output file atomically
            def _write_info(tmpf):
                tmpf.write('\n'.join(hidden_info).encode('utf-8'))
            
            self._atomic_write_file(output_path, _write_info)
            return True, f"Hidden information extracted to {output_path}"
            
        except Exception as e:
            return False, f"Hidden info extraction failed: {str(e)}"
    
    def pdf_to_word(self, input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert PDF to Word document
        
        Args:
            input_path: Path to input PDF file
            output_path: Path to output DOCX file
            
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            # Try to import python-docx
            try:
                from docx import Document
                from docx.shared import Inches
            except ImportError:
                return False, "python-docx not installed. Please install with: pip install python-docx"
            
            if not fitz:
                return False, "PyMuPDF not available for PDF to Word conversion"
            
            # Open PDF with PyMuPDF
            doc = Document()
            pdf_document = fitz.open(input_path)
            
            total_pages = len(pdf_document)
            for page_num in range(total_pages):
                self.update_progress(int((page_num + 1) / total_pages * 100))
                if self._cancel_requested:
                    return False, "Operation cancelled"
                
                page = pdf_document.load_page(page_num)
                text = page.get_text()
                
                # Add page content to document
                doc.add_heading(f'Page {page_num + 1}', level=1)
                doc.add_paragraph(text)
                
                # Try to extract images
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    # Save image temporarily using secure temp file
                    tmp_fd = None
                    tmp_name = None
                    try:
                        tmp_fd, tmp_name = tmp_mkstemp(suffix='.' + image_ext)
                        os.write(tmp_fd, image_bytes)
                        os.close(tmp_fd)
                        tmp_fd = None
                        doc.add_picture(tmp_name, width=Inches(4))
                    except Exception:
                        pass
                    finally:
                        if tmp_fd is not None:
                            try:
                                os.close(tmp_fd)
                            except Exception:
                                pass
                        if tmp_name and os_path.exists(tmp_name):
                            try:
                                os.remove(tmp_name)
                            except Exception:
                                pass
            
            pdf_document.close()
            
            # Save Word document atomically
            def _save_docx(tmp_path):
                doc.save(tmp_path)
            
            self._atomic_write_via_path(output_path, _save_docx)
            return True, f"PDF converted to Word document: {output_path}"
            
        except Exception as e:
            return False, f"PDF to Word conversion failed: {str(e)}"
    
    def _show_compression_error_popup(self):
        """
        Show a custom popup with compression error gif when no compression is achieved
        """
        try:
            if not tk or not Toplevel or not Image:
                return
            
            # Create popup window
            popup = Toplevel()
            popup.title("Compression Info")
            popup.geometry("550x400")
            popup.resizable(False, False)
            # Disable menu bar
            popup.overrideredirect(False)
            
            # Center the window
            popup.transient()
            popup.grab_set()
            
            # Load and display the gif using package-relative path
            module_dir = Path(__file__).parent
            gif_path = module_dir / "assets" / "compression_err.gif"
            if gif_path.exists():
                try:
                    # Load the GIF and handle animation
                    gif_image = Image.open(str(gif_path))
                    frames = []
                    
                    try:
                        # Extract all frames from the GIF
                        while True:
                            frames.append(ImageTk.PhotoImage(gif_image.copy()))
                            gif_image.seek(gif_image.tell() + 1)
                    except EOFError:
                        pass  # End of frames
                    
                    # Display animated GIF
                    img_label = Label(popup)
                    img_label.pack(pady=10)
                    
                    def animate_gif(frame_index=0):
                        if frames:
                            img_label.config(image=frames[frame_index])
                            popup.after(100, animate_gif, (frame_index + 1) % len(frames))
                    
                    animate_gif()
                    
                except Exception:
                    # If gif loading fails, show text instead
                    Label(popup, text="Compression Info", font=("Calibri", 16, "bold")).pack(pady=10)
                    self.logger.error("Error loading compression error GIF", exc_info=True)
            else:
                # If gif file doesn't exist, show icon
                Label(popup, text="Compression Info", font=("Calibri", 16, "bold")).pack(pady=10)
            
            # Info message with better formatting
            info_text = (
                "Compression completed but no size reduction detected.\n"
                "This file has already been optimized or contains\n"
                "elements that cannot be compressed further.\n"
                "If you need further compression, consider using the\n"
                "'Microsoft Print to PDF' option from the print dialog."
            )
            
            Label(popup, text=info_text, justify="center", wraplength=400, 
                  font=("Calibri", 10), padx=20, pady=10).pack(pady=10)
            
            # OK button
            Button(popup, text="OK", command=popup.destroy, width=10, 
                   font=("Calibri", 10)).pack(pady=15)
            
            # Center the popup on screen
            popup.update_idletasks()
            x = (popup.winfo_screenwidth() // 2) - (popup.winfo_width() // 2)
            y = (popup.winfo_screenheight() // 2) - (popup.winfo_height() // 2)
            popup.geometry(f"+{x}+{y}")
            
        except Exception:
            # Fallback to simple messagebox if custom popup fails
            if messagebox:
                messagebox.showinfo("Compression Info", 
                    "Compression completed but no size reduction detected.\n"
                    "This file has already been optimized or contains elements that cannot be compressed further.\n"
                    "If you need further compression, consider using the 'Microsoft Print to PDF' option from the print dialog.")