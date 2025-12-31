"""
PDF to Word/DOCX Conversion Module for SafePDF
Handles PDF to Word document conversion operations
"""

import os
from os import path as os_path
from tempfile import mkstemp as tmp_mkstemp
from typing import Tuple

import pypdfium2 as pdfium
from PIL import Image
from PyPDF2 import PdfReader

from SafePDF.logger.logging_config import get_logger


class PDFToWordConverter:
    """Class handling PDF to Word/DOCX conversion operations"""

    def __init__(self, progress_callback=None, language_manager=None, atomic_write_via_path=None):
        """
        Initialize PDF to Word converter

        Args:
            progress_callback: Function to call for progress updates (0-100)
            language_manager: Language manager for localization
            atomic_write_via_path: Function for atomic file writing via path
        """
        self.progress_callback = progress_callback
        self.language_manager = language_manager
        self._atomic_write_via_path = atomic_write_via_path
        self._cancel_requested = False
        self.logger = get_logger("SafePDF.PDF2Word")

    def update_progress(self, value):
        """Update progress if callback is available"""
        if self.progress_callback:
            self.progress_callback(value)

    def request_cancel(self):
        """Request cancellation of a running operation."""
        self._cancel_requested = True

    def pdf_to_word(self, input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert PDF to Word document using PyPDF2 for text and pypdfium2 for images

        Args:
            input_path: Path to input PDF file
            output_path: Path to output DOCX file

        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            # Try to import python-docx
            try:
                from docx import Document
                from docx.shared import Inches
            except ImportError:
                return (
                    False,
                    self.language_manager.get(
                        "op_docx_unavailable", "python-docx not installed. Please install with: pip install python-docx"
                    )
                    if self.language_manager
                    else "python-docx not installed. Please install with: pip install python-docx",
                )

            if not PdfReader:
                return False, self.language_manager.get(
                    "op_pypdf_unavailable", "PyPDF2/pypdf not available"
                ) if self.language_manager else "PyPDF2/pypdf not available"

            # Open PDF with PyPDF2
            doc = Document()

            with open(input_path, "rb") as file:
                reader = PdfReader(file)
                total_pages = len(reader.pages)

                for page_num in range(total_pages):
                    self.update_progress(int((page_num + 1) / total_pages * 100))
                    if self._cancel_requested:
                        return False, self.language_manager.get(
                            "op_word_cancelled", "Operation cancelled"
                        ) if self.language_manager else "Operation cancelled"

                    page = reader.pages[page_num]
                    text = page.extract_text()

                    # Add page content to document
                    doc.add_heading(f"Page {page_num + 1}", level=1)
                    if text.strip():
                        doc.add_paragraph(text)
                    else:
                        doc.add_paragraph("[No text content detected on this page]")

                    # Try to extract images using pypdfium2 if available
                    if pdfium and Image:
                        try:
                            # Open PDF and render single page
                            pdf = pdfium.PdfDocument(input_path)
                            if page_num < len(pdf):
                                page = pdf[page_num]
                                pil_image = page.render(scale=2.0).to_pil()

                                # Save image temporarily
                                tmp_fd = None
                                tmp_name = None
                                try:
                                    tmp_fd, tmp_name = tmp_mkstemp(suffix=".png")
                                    os.close(tmp_fd)
                                    tmp_fd = None
                                    pil_image.save(tmp_name, "PNG")
                                    doc.add_picture(tmp_name, width=Inches(6))
                                except Exception:
                                    pass
                                finally:
                                    if tmp_fd is not None:
                                        try:
                                            os.close(tmp_fd)
                                        except Exception:
                                            pass
                                    if tmp_name and os_path.exists(tmp_name):
                                        try:
                                            os.remove(tmp_name)
                                        except Exception:
                                            pass
                            pdf.close()
                        except Exception:
                            pass  # Continue without images if conversion fails

            # Save Word document atomically
            def _save_docx(tmp_path):
                doc.save(tmp_path)

            if self._atomic_write_via_path:
                self._atomic_write_via_path(output_path, _save_docx)
            else:
                doc.save(output_path)

            success_msg = (
                self.language_manager.get("op_word_success", "PDF converted to Word document: {output_path}")
                if self.language_manager
                else "PDF converted to Word document: {output_path}"
            )
            return True, success_msg.format(output_path=output_path)

        except Exception as e:
            error_msg = (
                self.language_manager.get("op_word_failed", "PDF to Word conversion failed: {error}")
                if self.language_manager
                else "PDF to Word conversion failed: {error}"
            )
            return False, error_msg.format(error=str(e))
